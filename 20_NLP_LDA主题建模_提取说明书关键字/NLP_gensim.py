# -*- coding: utf-8 -*-
"""
Created on Tue Mar 20 19:10:03 2018

@author:richieBao-caDesign设计(cadesign.cn)
"""

'''
主题模型（Topic Model）在机器学习和自然语言处理等领域是用来在一系列文档中发现抽象主题的一种统计模型。
主题模型自动分析每个文档，统计文档内的词语，根据统计的信息来断定当前文档含有哪些主题，以及每个主题所占的比例各为多少。
'''
from gensim import corpora, models, similarities #调入gensim的语料库、模型库和相似度比较库
from collections import defaultdict
import docx
import jieba #分词中文处理库
import os
import numpy as np
from pprint import pprint
import pandas as pd

'''gensim使用python标注的日志类来记录不同优先级的各种事件，使用如下语句激活日志，从而运行时亦可查看。'''
import logging
logging.basicConfig(format='%(asctime)s : %(levelname)s : %(message)s', level=logging.INFO)


'''A-语料库与向量空间 Corpora and Vector Spaces'''
'''01-a-读取.docx文档，并将文本内容按照段落存储于定义的documents列表中。'''

#fn_a=r'E:\MUBENAcademy\pythonSystem\code\xianmasterplanning.docx'
fn_a=r'E:\MUBENAcademy\pythonSystem\code\digitalDesignFinished12182013.docx'
file_a=docx.Document(fn_a)
print("段落数:"+str(len(file_a.paragraphs)))
documents=[]
for para in file_a.paragraphs:
    documents.append(para.text)
    
#for i in range(len(file.paragraphs)):
#    print(file.paragraphs[i].text)
#someOfText=file.paragraphs[10].text
#seg=jieba.cut(someOfText)
#print(','.join(seg))
#seg=jieba.cut()

'''01-b-读取.xlsx文件。'''
#fn_b=r'E:\MUBENAcademy\pythonSystem\code\data\hornetsnestData.xlsx' #使用八爪鱼抓取马蜂窝部分游记数据。
#file_b=pd.read_excel(fn_b,sheetname=0,header=0,usecols=[4])
#documents=file_b.values.tolist()
#flatten_lst=lambda lst: [m for n_lst in lst for m in flatten_lst(n_lst)] if type(lst) is list else [lst]
#documents=flatten_lst(documents)

'''02-读取停用词文档，并将停用词存储于定义的stopset集合中。并分词以及移除语料库中的停用词'''    
#stopwordsfn=r'E:\MUBENAcademy\pythonSystem\code\stopwords.docx'   
#stopwords=docx.Document(stopwordsfn)
#stoplist=[]
#for sw in stopwords.paragraphs:
#    stoplist.append(sw.text)
stopwords=r'E:\MUBENAcademy\pythonSystem\code\stopwords.txt' 
stoplist=[]
print(documents[11])
f=open(stopwords,'r')
try:
    for line in f:
        stoplist.append(line)   
finally:
    f.close()   
stopset=set(stoplist)

texts=[[word for word in jieba.cut(document) if word not in stopset] for document in documents if len(document)!=0]
texts = [[word for word in text if word!=u' ' and len(word)>1] for text in texts] #对于有些文章，去除多余空格，并限制为词组，而不是单个字。

'''03-去除仅出现1次的单词'''
frequency = defaultdict(int)
for text in texts:
    for token in text:
        frequency[token] += 1
texts = [[token for token in text if frequency[token] > 1] for text in texts] 
   
#pprint(texts)

'''04-利用corpora.Dictionary()为每一个出现在语料库中的单词/词汇分配独一无二的整数编号。'''
dictionary = corpora.Dictionary(texts)
dictionary.save('/tmp/NLPexperi.dict') #并把字典存储到硬盘中，方便日后调用。
#print(dictionary)
#print(dictionary.token2id)
#new_doc="历史沿革：公元前11世纪，周文王在沣河西岸建立丰京，武王在沣河东岸建立镐京，统一华夏，丰镐两京隔河相望，成为世界上第一个“双子城”，也开创了西安长期作为中国古代政治、经济、文化中心的历史；秦统一全国后，在今西安北郊建了许多著名的宫殿，如阿房宫、章台宫、兴乐宫、信宫等,地方上实行郡县两级行政建制，设内史辖京畿各县(内史政区与官职同名，为郡级建制)，今西安市辖域属其管辖范围。西汉都长安,筑长安城墙，城区面积36平方公里，主要宫殿位于城的南部；隋唐两代均都长安，隋称大兴城，唐改名长安城，在长安城周围的京畿地区，隋唐均设京兆尹(郡、府)或雍州，作为郡级建制以统长安、大兴(唐改为万年)等20余县，唐以后，长安城不复为都，发展受到一定影响，但仍不失为一个重要的地方性都会；明初设西安府，“西安”由此得名，明将五代、宋、元的长安城向北、向东加以扩展，形成今天西安古城的规模；清继续沿用西安府名，辖长安、咸宁等15县，1州，孝义(今柞水)、宁陕(今县)等2厅；1913年，民国政府在关中地区设立关中道，辖长安等41县；1947年，西安升为全国13个直辖市之一；1949年，中国人民解放军成立西安市人民政府，辖12个区，隶属于陕甘宁边区政府。1953年升为中央直辖市，1954年又改为省辖市，同时把原来12个区调整为9个(即新城区、碑林区、莲湖区、灞桥区、未央区、雁塔区，再加上长乐区、阿房区和草滩区)。从此以后西安市的领属关系再未发生变化，但所辖区县不断变迁，直到1983年才稳定下来。"
#new_vec=dictionary.doc2bow(jieba.cut(new_doc))
#print(new_vec)
'''05-建立BOW稀疏文档向量。'''
corpus = [dictionary.doc2bow(text) for text in texts]
corpora.MmCorpus.serialize('/tmp/NLPexperi.mm', corpus) #将BOW稀疏向量存入硬盘，方便日后调用。
#print(corpus)


'''B-主题与转换 Topics and Transformations'''
'''06-调入存入硬盘的语料库和BOW稀疏向量'''
if (os.path.exists("/tmp/NLPexperi.dict")):
    dictionary = corpora.Dictionary.load('/tmp/NLPexperi.dict')
    corpus = corpora.MmCorpus('/tmp/NLPexperi.mm')
    print("loaded.~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~")
else:
    print("fail loading.")
'''
07-初始化模型(一种向量表示方式转换为另一种)，并对整个语料库实施转换。
转换的目的：
1.将语料库中隐藏的结构发掘出来，发现词语之间的关系，并且利用这些结构、关系使用一种新的、更有语义价值的（这是我们最希望的）方式描述其中的文档。
2.使得表示方式更加简洁。不仅能提高效率（新的表示方法一般消耗较少的资源）同时提高效果（忽略边际数据趋势、降低噪音）。
'''
tfidf = models.TfidfModel(corpus) #词频-逆向文件频率（tf-idf）的向量化参数
corpus_tfidf = tfidf[corpus]
#for doc in corpus_tfidf:
#    print(doc)

'''08-用隐含狄利克雷分布(Latent Dirichlet Allocation,LDA)主题建模'''
ldamodel=models.ldamodel.LdaModel(corpus,id2word=dictionary,num_topics=20,passes=25)
ldamodel.save('/tmp/ldamodel.lda') #保持模型
ldamodel=models.LdaModel.load('/tmp/ldamodel.lda')
ldamodel.print_topics()

#topics_matrix=ldamodel.show_topics(formatted=False,num_words=20)
#topics_matrix=np.array([[j[0] for j in i[1]] for i in topics_matrix])
#for i in topics_matrix:
#    print([word for word in i],chr(hex))

'''利用潜在语义索引（LSI） 将tfidf语料转化为一个潜在2-D空间(num_topics=2)。并存储到硬盘空间。'''
#lsi = models.LsiModel(corpus_tfidf, id2word=dictionary, num_topics=200) #初始化一个LSI转换。num_topics推荐200-500
#corpus_lsi = lsi[corpus_tfidf] #在原始语料库上加上双重包装: bow->tfidf->fold-in-lsi
#print("~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~")
#lsi.print_topics(2)
##for doc in corpus_lsi:
##    print(doc)
#lsi.save('/tmp/model.lsi')  
#lsi = models.LsiModel.load('/tmp/model.lsi') 